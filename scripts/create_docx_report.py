import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Inches, Pt, RGBColor


def set_cell_background(cell, fill_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = OxmlElement("w:tcMar")
    for m, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        node = OxmlElement(f"w:{m}")
        node.set(qn("w:w"), str(val))
        node.set(qn("w:type"), "dxa")
        tc_mar.append(node)
    tc_pr.append(tc_mar)


def add_heading_styled(doc, text, level):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.keep_with_next = True

    run = p.add_run(text)
    run.font.bold = True

    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = RGBColor(15, 23, 42)
    elif level == 2:
        run.font.size = Pt(14)
        run.font.color.rgb = RGBColor(22, 119, 255)
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(51, 65, 85)
    return p


def create_report():
    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    print("Generating Cover Page...")
    # COVER PAGE
    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_school = p_top.add_run(
        "TRƯỜNG ĐẠI HỌC VINUNIVERSITY\nKHOA KỸ THUẬT & KHOA HỌC MÁY TÍNH\n-----------------------------------"
    )
    r_school.font.size = Pt(12)
    r_school.font.bold = True
    r_school.font.color.rgb = RGBColor(71, 85, 105)

    p_space = doc.add_paragraph()
    p_space.paragraph_format.space_before = Pt(40)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("BÁO CÁO NHÓM - THÀNH VIÊN C\nCHUYÊN MÔN: PRODUCT DESIGN & UI/UX LEAD")
    r_title.font.size = Pt(22)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(22, 119, 255)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_before = Pt(12)
    r_sub = p_sub.add_run(
        "DỰ ÁN: RidePulse DQ – Autonomous Data Quality & Anomaly Intelligence Platform\nMÔN HỌC: Product Development (P-028)"
    )
    r_sub.font.size = Pt(14)
    r_sub.font.italic = True
    r_sub.font.color.rgb = RGBColor(51, 65, 85)

    p_box = doc.add_paragraph()
    p_box.paragraph_format.space_before = Pt(120)
    p_box.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_info = p_box.add_run(
        "Thực hiện bởi: Thành viên C (UI/UX Lead)\nNhiệm vụ chính: Mục 3 - Wireframe & UI Flow (Figma & Ant Design)\nLớp: Product Development 2026\nNgày nộp: 31/07/2026"
    )
    r_info.font.size = Pt(11)
    r_info.font.color.rgb = RGBColor(15, 23, 42)

    doc.add_page_break()

    # Setup Header
    section = doc.sections[0]
    header = section.header
    hp = header.paragraphs[0]
    hp.text = "RidePulse DQ — Báo cáo Công việc Thành viên C (UI/UX Lead)"
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hp.runs[0].font.size = Pt(9)
    hp.runs[0].font.color.rgb = RGBColor(148, 163, 184)

    # CHAPTER 1
    add_heading_styled(doc, "CHƯƠNG 1: TỔNG QUAN DỰ ÁN & BÀI TOÁN NGHIỆP VỤ", 1)

    p = doc.add_paragraph()
    p.add_run("Trong dự án ").font.size = Pt(11)
    r_b = p.add_run("RidePulse DQ – Autonomous Data Quality & Anomaly Intelligence Platform")
    r_b.font.bold = True
    p.add_run(", ").font.size = Pt(11)
    r_role = p.add_run("Thành viên C")
    r_role.font.bold = True
    r_role.font.color.rgb = RGBColor(22, 119, 255)
    p.add_run(" chịu trách nhiệm toàn bộ ").font.size = Pt(11)
    p.add_run(
        "Mục 3: Wireframe & UI Flow, thiết kế cấu trúc luồng thao tác người dùng, phát thảo giao diện Ant Design và lập trình bản Web Application Interactive Prototype."
    ).font.size = Pt(11)

    add_heading_styled(doc, "1.1 Bối cảnh Dữ liệu Vận hành Ride-Hailing", 2)
    p = doc.add_paragraph()
    p.add_run("Hệ thống xử lý 4 tập dữ liệu vận hành chính: ").font.size = Pt(11)
    p.add_run(
        "dich_vu_xe_trips (chuyến đi), dich_vu_xe_drivers (tài xế), dich_vu_xe_customers (hành khách), và dich_vu_xe_payments (thanh toán). "
    ).font.bold = True
    p.add_run(
        "Dữ liệu thường xuyên mắc phải các lỗi dữ liệu bẩn (Bad Data) như: NULL ở khóa chính, cước âm (fare_amount < 0), vi phạm khoảng giá trị (outliers), sai format timestamp, và freshness lag. Giải pháp AI Agent + HITL giúp tự động hóa khâu đọc metadata, đề xuất rule, kiểm thử dbt và phát hiện bất thường bằng ML."
    ).font.size = Pt(11)

    # CHAPTER 2
    add_heading_styled(doc, "CHƯƠNG 2: PHÂN TÍCH VAI TRÒ NGƯỜI DÙNG & PHÂN QUYỀN", 1)
    p = doc.add_paragraph()
    p.add_run("1. Data Steward (Core Operator): ").font.bold = True
    p.add_run(
        "Có toàn quyền connect dataset, chạy profiling, duyệt AI rules (Approve/Reject/Edit), thực thi dbt test và xem chẩn đoán nguyên nhân gốc AI Root Cause Diagnosis.\n"
    ).font.size = Pt(11)
    p.add_run("2. Viewer (Executive Lead): ").font.bold = True
    p.add_run(
        "Giao diện Read-Only an toàn, chỉ xem Dashboard Data Health Score, các sự cố và báo cáo Trend Analysis mà không thể thao tác sửa đổi quy tắc."
    ).font.size = Pt(11)

    # CHAPTER 3 (MỤC 3 CỦA YÊU CẦU ĐỀ BÀI)
    add_heading_styled(doc, "MỤC 3: WIREFRAME & UI FLOW (THÀNH VIÊN C ĐẢM NHẬN)", 1)

    add_heading_styled(doc, "A. UI Flow (Sơ đồ luồng đi người dùng)", 2)
    p = doc.add_paragraph()
    p.add_run("Luồng thao tác được thiết kế mạch lạch qua 7 bước theo đúng yêu cầu đề bài:\n").font.size = Pt(11)
    p.add_run(
        "1. Đăng nhập (Login): Chọn vai trò Steward hoặc Viewer.\n"
        "2. Dashboard Tổng quan: Xem chỉ số Data Health Score chung (87.4%).\n"
        "3. Select Dataset: Chọn bảng dữ liệu vận hành dich_vu_xe_trips.\n"
        "4. Agent Profiling & Rule Proposal: AI hiện gợi ý các Rule (Not-null, Range check, Format).\n"
        "5. HITL Review (Chỉ Steward): Checkbox Duyệt / Sửa / Từ chối Rule trực quan.\n"
        "6. Test Execution & Anomaly Detection: Hệ thống chạy test dbt & mô hình ML Isolation Forest.\n"
        "7. Detail Report / Alerting & Trend Analysis: Xem chi tiết lỗi, chẩn đoán nguyên nhân AI Diagnosis và biểu đồ xu hướng."
    ).font.size = Pt(11)

    add_heading_styled(doc, "B. Khung Màn Hình Chính (Wireframe Layout - Ant Design Style)", 2)
    p = doc.add_paragraph()
    p.add_run(
        "Thành viên C phác thảo các màn hình trọng tâm theo yêu cầu với phong cách Ant Design Admin Dashboard (Frame 1440px Desktop Grid):\n"
    ).font.size = Pt(11)
    p.add_run(
        "• Màn 1: Rule Review Screen (HITL): Bảng chứa danh sách Rule do AI tạo (Column Name, Rule Type, AI Reason, Confidence %, Suggested Threshold, Status: Pending/Approved/Rejected, Action buttons).\n"
        "• Màn 2: Anomaly Dashboard: Biểu đồ Time-series hiển thị các điểm bất thường (Anomaly dots màu đỏ), bảng danh sách Alert đi kèm nút '🤖 AI Diagnosis' (nhấn vào hiện Modal giải thích nguyên nhân gốc rễ).\n"
        "• Màn 3: Trend & Evaluation Screen: Biểu đồ đường thể hiện chỉ số Data Quality Score theo tuần/tháng và các thông số Precision (94.2%), Recall (91.8%), F1-Score (93.0%) của mô hình ML."
    ).font.size = Pt(11)

    # CHAPTER 4: IMAGES
    add_heading_styled(doc, "CHƯƠNG 4: HỒ SƠ HÌNH ẢNH UI PROTOTYPE THỰC TẾ (11 SCREENS)", 1)

    img_dir = os.path.abspath("docs/images")
    screen_docs = [
        (
            "Screen 1: Đăng nhập & Chọn Phân quyền (Login & Role Selection)",
            "screen_1_login.png",
            "Màn hình đăng nhập chọn vai trò Data Steward hoặc Viewer.",
        ),
        (
            "Screen 2: Steward Dashboard (Data Health Score 87.4%)",
            "screen_2_steward_dashboard.png",
            "Dashboard tổng quan với điểm số Data Health Score 87.4%.",
        ),
        (
            "Screen 3: Catalog Lựa chọn Dataset (dich_vu_xe_trips)",
            "screen_3_dataset_catalog.png",
            "Danh mục tra cứu bảng dữ liệu vận hành gọi xe.",
        ),
        (
            "Screen 4: Dataset Profiling & Metadata Insights",
            "screen_4_dataset_profiling.png",
            "Kết quả AI Profiling phân tích Null %, Unique % và Outliers.",
        ),
        (
            "Screen 5: Màn 1 - Rule Review Screen (HITL Review Table)",
            "screen_5_ai_rule_proposals.png",
            "Bảng HITL review chứa danh sách Rule do AI gợi ý với các nút Approve/Reject/Edit.",
        ),
        (
            "Screen 6: Màn 1 (Phụ) - Rule Edit Modal",
            "screen_6_rule_edit_modal.png",
            "Modal chỉnh sửa thông số Threshold, Severity và Mô tả rule.",
        ),
        (
            "Screen 7: Running Tests & Streaming Console Log",
            "screen_7_execution_log.png",
            "Chạy dbt test suite với Stepper 4 bước và Live Terminal Log.",
        ),
        (
            "Screen 8: Màn 2 - Anomaly Dashboard & Alert Stream",
            "screen_8_anomaly_dashboard.png",
            "Biểu đồ Time-Series gắn đốm đỏ Anomaly dots và bảng Alert có nút AI Diagnosis.",
        ),
        (
            "Screen 9: Màn 2 (Phụ) - AI Diagnosis Modal",
            "screen_9_ai_diagnosis_modal.png",
            "Modal giải thích nguyên nhân gốc do AI Agent chẩn đoán.",
        ),
        (
            "Screen 10: Màn 3 - Trend & Evaluation Screen",
            "screen_10_trend_analysis.png",
            "Biểu đồ xu hướng 30 ngày và các chỉ số ML Metrics (Precision, Recall, F1).",
        ),
        (
            "Screen 11: Executive Viewer Dashboard (Read-Only View)",
            "screen_11_viewer_dashboard.png",
            "Giao diện Read-Only an toàn 100% dành cho Viewer.",
        ),
    ]

    for title, img_name, desc in screen_docs:
        add_heading_styled(doc, title, 2)
        img_path = os.path.join(img_dir, img_name)
        if os.path.exists(img_path):
            p_img = doc.add_paragraph()
            p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_img = p_img.add_run()
            run_img.add_picture(img_path, width=Inches(5.8))

            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r_cap = p_cap.add_run(f"Hình: {title}")
            r_cap.font.size = Pt(9.5)
            r_cap.font.italic = True
            r_cap.font.color.rgb = RGBColor(100, 116, 139)

        p_desc = doc.add_paragraph()
        p_desc.paragraph_format.space_before = Pt(4)
        p_desc.paragraph_format.space_after = Pt(12)
        r_d = p_desc.add_run(f"📌 Phân tích UX/UI: {desc}")
        r_d.font.size = Pt(10.5)
        r_d.font.color.rgb = RGBColor(51, 65, 85)

    doc.add_page_break()

    # CHAPTER 5 & 6
    add_heading_styled(doc, "CHƯƠNG 5: TỔNG KẾT BÀN GIAO SẢN PHẨM CỦA THÀNH VIÊN C", 1)
    p_final = doc.add_paragraph()
    p_final.add_run(
        "Thành viên C bàn giao đầy đủ sản phẩm đáp ứng 100% yêu cầu đề bài:\n\n"
        "1. File Báo cáo Word (.docx 10 trang): docs/BAO_CAO_THANH_VIEN_C_UI_UX.docx\n"
        "2. Thư mục 11 Ảnh chụp UI chất lượng cao: docs/images/\n"
        "3. Bản Web App Prototype tương tác: ui_test/index.html\n"
        "4. Tài liệu Wireframe Specs: ridepulse_dq_design_spec.md"
    ).font.size = Pt(11)

    output_docx = os.path.abspath("docs/BAO_CAO_THANH_VIEN_C_UI_UX.docx")
    doc.save(output_docx)
    print(f"Word document updated successfully at: {output_docx}")


if __name__ == "__main__":
    create_report()
